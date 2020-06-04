from docx import Document
from docx.shared import Inches

file = 'F:\\测试报告模板.docx'
fig_path = 'E:\\ProgramData\\PycharmProjects\\webtest\\Image\\2020_05_14_11_15_20.png'
fig_path2 = 'E:\\ProgramData\\PycharmProjects\\webtest\\Image\\2020_05_14_11_12_30.png'

doc = Document()
doc.add_heading(u'测试记录', level=1)
doc.add_paragraph(u'步骤1：XXXXXX')
doc.add_picture(fig_path, width=Inches(6.0), height=Inches(3.0))
doc.add_paragraph(u'步骤2：XXXXXX')
doc.add_picture(fig_path2, width=Inches(6.0), height=Inches(3.0))
doc.save('F:\\test_pic01.docx')


# document = Document(file)
# steptable = document.tables[0].cell(0, 1)
# # 在第一个表格的指定的单元格（第一行第一列）插入图片
# run = steptable.paragraphs[0].add_run()
# # run.add_paragraph('步骤1：打开菜单')
# steptable.add_paragraph('步骤1：打开菜单')
# steptable.add_paragraph('')
# run.add_picture(fig_path, width=Inches(6.0), height=Inches(3.0))
# steptable.add_paragraph('')
#
# # run.add_paragraph('步骤2： 查询XXXX')
# steptable.add_paragraph(u'步骤2：输入号码')
# run.add_picture(fig_path2, width=Inches(6.0), height=Inches(3.0))
# document.save('F:\\testReport01.docx')
