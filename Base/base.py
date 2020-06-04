from selenium.webdriver.common.action_chains import ActionChains  # 处理鼠标事件
from selenium.webdriver.support.ui import WebDriverWait  # 用于处理元素等待
from Common.Mylog import LogManager
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import *  # 用于处理异常
from selenium.webdriver.support.select import Select
from selenium.webdriver.common.keys import Keys
from Common import ReadConfig
import time,os,sys
from docx import Document
from docx.shared import Inches

logger = LogManager('base').get_logger_and_add_handlers(1,is_add_stream_handler=True, log_path=ReadConfig.log_path, log_filename=time.strftime("%Y-%m-%d")+'.log')

class Base():
    def __init__(self,driver):
        self.driver = driver
        self.timeout = 5
        self.t = 1
        self.doc = Document()

    def return_driver(self):
        return self.driver
     #id
    def byid(self,element):
        return self.driver.find_element_by_id(element)
    #name
    def byname(self,element):
        return self.driver.find_element_by_name(element)
    #xpath
    def byxpath(self,element):
        return  self.driver.find_element_by_xpath(element)
    #css
    def bycss(self,element):
        return  self.driver.find_element_by_css_selector(element)
    #class
    def byclassname(self,element):
        return self.driver.find_element_by_class_name(element)
    #url
    def dr_url(self):
        return self.driver.current_url

    def findele(self,*args):
        try:
            logger.info("通过"+args[0]+"定位，元素是"+args[1])
            return self.driver.find_element(*args)
        except:
            logger.info("定位元素失败！")
            self.screen()

    def find_element(self, locator):
        """定位元素，参数locator为原则,利用显式等待"""
        try:
            logger.info("通过"+locator[0]+"定位，元素是"+locator[1])
            return self.driver.find_element(*locator)
        except:
            logger.info("页面未找到元素:"+str(locator))
            self.screen()

    def findeles(self,*locator):
        '''更加xpath-str获取多个元素'''
        try:
            print(*locator)
            logger.info("通过" + locator[0] + "定位，元素是" + locator[1])
            return self.driver.find_elements(*locator)
        except:
            logger.info("页面未找到元素:"+str(locator))
            self.screen()
            return False

    def find(self, locator,value = ''):
        """定位到元素，返回元素对象，没定位到，Timeout异常 loctor 传元祖，如（"id", "kw"）
        eg : input_usr = find(("id","username"))
        """
        if not isinstance(locator, tuple):
            logger.info('locator参数类型错误，必须传元祖类型：loc = ("id", "value1")')
        else:
            logger.info("正在定位元素信息：定位方式" + locator[0] + "  元素值->" + locator[1] + "，value值->" + str(value))
        if value == '':  # 默认用这种方式
            ele = WebDriverWait(self.driver, self.timeout, self.t).until(EC.presence_of_element_located(locator))
            if ele:
                logger.info("定位到该元素{}" .format(str(ele)))
                return ele
            else:
                logger.info("定位失败：定位方式->" + locator[0] + ", value值->" + locator[1])
                self.screen()    #找不到就自动截屏
                return False
        else:  #value值定位
            ele = WebDriverWait(self.driver, self.timeout, self.t).until(
                EC.text_to_be_present_in_element_value(locator, value))
            logger.info('定位到该元素:{}'.format(str(ele)))
            return ele

    def finds(self, locator, value=''):
        '''定位一组元素，返回元素对象list，没定位到，Timeout异常 '''
        if value == '':  #默认为此常规定位方法,只要找到一个，返回true
            # eles = WebDriverWait(self.driver, self.timeout, self.t).until(EC.presence_of_all_elements_located(locator))
            eles = self.find_eles(locator)
            if eles:
                return eles
            else:
                print("定位失败：定位方式->{locator[0]}, value值->{locator[1]}")
                logger.info("定位失败：定位方式->{locator[0]}, value值->{locator[1]}")
                self.screen()  #保存截图
                return []
        else:  # value值定位
            eles = WebDriverWait(self.driver, self.timeout, self.t).until(
                EC.presence_of_element_located(locator, value))
            return eles


    '''=======================判断元素方法收集========================'''
    def isElementExist(self, locator,Type=''):
        """判断单个元素是否在DOM里面 （是否存在）"""
        try:
            ele = self.find(locator)
            if Type =='':
                # self.find(locator)
                return True
            elif Type == 'click':
                ele.click()
                time.sleep(2)
            elif Type == 'text':
                return self.get(locator,Type='text')
            else:
                print("Type参数%s错误，仅可为click或者不传或者text", Type)
        except:
            return False

    def isElementExists(self, locator):
        ''' 判断一组元素是否在DOM里面 （是否存在），若不存在，返回一个空的list'''
        eles = self.finds(locator)
        n = len(eles)
        if n == 0:
            return False
        elif n == 1:
            return True
        else:
            print("定位到元素的个数：",n)
            logger.info("定位到元素的个数:" + str(n))
            return True

    def isElementDisplay(self, locator,action =''):
        """判断单个元素是否显示在页面上"""
        try:
            ele = self.find(locator)
            if action =='':
                r = ele.is_displayed()
                return r
            elif action == 'click':  # 如果type参数为click，执行元素的点击操作
                ele.click()
                time.sleep(2)
            elif action == 'text':
                return self.get(locator,Type='text')
            else:
                print("action参数%s错误，仅可为click、text或者不传", action)
        except:
            return False

    def isSelected(self, locator, Type=''):
        ''' 判断元素是否被选中，返回bool值 及点（选中/取消选中）'''
        try:
            ele = self.find(locator)
            if Type == '':  # 如果type参数为空，返回元素是否为选中状态，True/False (默认)
                r = ele.is_selected()
                return r
            elif Type == 'click':  # 如果type参数为click，执行元素的点击操作
                ele.click()
            else:
                print("type参数%s错误，仅可为click或者不传",Type)
        except:
            return False

    def is_text_in_element(self, text, locator):
        u"""判断是否定位到元素"""
        try:
            WebDriverWait(self.driver, self.timeout,self.t).until(EC.presence_of_element_located(locator,text))
            return True
        except TimeoutException:
            print("元素未定位到:" + str(locator))
            logger.info("元素未定位到:"+str(locator))
            return False

    def is_element_located(self, locator):
        u"""判断是否定位到元素"""
        try:
            WebDriverWait(self.driver, self.timeout, 1).until(EC.presence_of_element_located(locator))
            return True
        except TimeoutException:
            print(u"元素未定位到:" + str(locator))
            logger.info(u"元素未定位到:" + str(locator))
            return False

    def is_element_displayed(self,element):
        '''传入元素，判断是否显示'''
        flag = element.is_displayed()
        return flag

    '''=======================判断元素方法收集========================'''

    def click(self,loactor):
        self.find(loactor).click()

    #输入值
    def sendkey(self,locator,value):
        ele = self.find(locator)
        try:
            self.js_scrollIntoView(ele)
            ele.click()
            ele.clear() #先清理一下数据
            ele.send_keys(value)
        except:
            logger.info("元素定位%s出现异常，无法设置value",locator)

    def clear(self,locator):
        ele = self.find(locator)
        try:
            ele.clear()  #清理一下
        except:
            logger.info("元素定位%s出现异常，无法清理数据",locator)

    def select(self, locator, value):
        """
        构造函数。对给定的元素进行了检查，确实是一个SELECT标记。如果不是,
        然后抛出一个意料之外的tag name exception.
        :Args:
         - css - element SELECT element to wrap
         - value - The value to match against
        Usage:
            <select name="NR" id="nr">
                <option value="10" selected="">每页显示10条</option>
                <option value="20">每页显示20条</option>
                <option value="50">每页显示50条</option>
            </select>
            driver.select("#nr", '20')
            driver.select("xpath=>//[@name='NR']", '20')
        """
        el = self.find(locator)
        Select(el).select_by_value(value)

    ##js处理
    """==============================js与jQuery相关====================================="""
    def js(self,str):
        self.driver.execute_script(str)

    def js_focus_element(self, locator):
        """聚焦元素 """
        target = self.find(locator)
        self.driver.execute_script("arguments[0].scrollIntoView();", target)


    def js_scroll_top(self):
        ''' 滚动到顶部 '''
        js = "window.scrollTo(0,0)"
        self.driver.execute_script(js)

    def js_scroll_end(self):
        ''' 滚动到底部 '''
        js = "window.scrollTo(0,document.body.scrollHeight)"
        self.driver.execute_script(js)

    def js_scrollIntoView(self,loc):
        '''鼠标移动到当前元素位置'''
        self.driver.execute_script("arguments[0].scrollIntoView();", loc)

    def js_findById(self, IdElement,action):
        ''' js查找元素，并做相应操作（传入id属性） 输入值：value='XXX' 点击：onclick '''
        js = "document.getElementById('%s').%s" %(IdElement,action)
        self.driver.execute_script(js)

    def js_findEle(self,Type,element, action):
        ''' js查找元素，并做相应操作 输入值：value='XXX' 点击：click()
        js定位仅可为：id、Name、TagName、ClassName、Selector（CSS） '''
        list = ['Name', 'TagName', 'ClassName', 'Selector']
        if type in list:
            # print("正在执行js操作：定位方式->{Type}, 元素值->{element}， 执行操作->{action}")
            print("正在执行js操作：定位方式->%s, 元素值->%s， 执行操作->%s") % (Type, element,  action)
            if type == 'Selector':
                # js = 'document.query{Type}All("{element}").{action}'
                js = 'document.query%s("%s").%s' % (Type, element, action)
                self.driver.execute_script(js)
            else:
                js = 'document.getElementsBy%s(%s).%s;' % (Type, element, action)
            self.driver.execute_script(js)
        else:
            print("type参数 %s 错误，js定位仅可为：'Name'、'TagName'、'ClassName'、'Selector'（CSS）" % Type)


    def js_finds(self, Type, element, index, action):
        ''' js查找元素，并做相应操作 输入值：value='XXX' 点击：click()
        js定位仅可为：id、Name、TagName、ClassName、Selector（CSS） '''
        list = ['Name', 'TagName', 'ClassName', 'Selector']
        if type in list:
            # print("正在执行js操作：定位方式->{Type}, 元素值->{element}， 下标值->{index}， 执行操作->{action}")
            print("正在执行js操作：定位方式->%s, 元素值->%s， 下标值->%s， 执行操作->%s") %(Type,element,index,action)
            if type == 'Selector':
                # js = 'document.query{Type}All("{element}"){index}.{action}'
                js = 'document.query%s("%s")%s.%s' %(Type,element,index,action)
            else:
                js = 'document.getElementsBy%s(%s)[%s].%s;' %(Type,element,index,action)
            self.driver.execute_script(js)
        else:
            print("type参数 %s 错误，js定位仅可为：'Name'、'TagName'、'ClassName'、'Selector'（CSS）" %Type)

    def js_readonly(self, idElement, value):
        ''' 去掉只读属性，并输入内容 一般为id '''
        # js = 'document.getElementById({idElement}).removeAttribute("readonly");document.getElementById({idElement}).value="{value}"'
        js = 'document.getElementById("%s").removeAttribute("readonly");document.getElementById("%s").value="%s"' %(idElement,idElement,value)
        self.driver.execute_script(js)

    def js_removeNullable(self, idElement):
        ''' 删除nullable属性 '''
        # js = 'document.getElementById({idElement}).removeAttribute("readonly");document.getElementById({idElement}).value="{value}"'
        js = 'document.getElementById("%s").removeAttribute("nullable");' % idElement
        self.driver.execute_script(js)

    def js_displayed(self, idElement,name):
        ''' 去掉只读属性，并输入内容 一般为id '''
        # js = 'document.getElementById({idElement}).removeAttribute("readonly");document.getElementById({idElement}).value="{value}"'
        # js = 'document.getElementById({idElement}).{name}.display'
        js = 'document.getElementById("%s").%s.display' %(idElement,name)
        self.driver.execute_script(js)

    def js_setsetAttribute(self,IdElement,name,value):
        '''通过js设置attribute属性值'''
        js = 'document.getElementById("%s").setAttribute("%s","%s")' %(IdElement,name,value)
        self.driver.execute_script(js)

    def js_setText(self,locator,val):
        target = self.find(locator)
        js = "arguments[0].innerText='%s';" %val
        self.driver.execute_script(js, target)

    def js_iframe(self, Type, element, action, index=''):
        ''' Js处理iframe 无需先切换到iframe上，再切回来操作
        输入值：value='' 点击：click() type=id时，index='' '''
        # js = 'document.getElementBy{Type}({element}){index}.contentWindow.document.body.{action}'
        js = 'document.getElementBy%s(%s)%s.contentWindow.document.body.%s' %(Type,element,index,action)
        self.driver.execute_script(js)

    def js_winstop(self):
        return self.driver.execute_script('window.stop()')

    """==============================js与jQuery相关====================================="""
    def switch_alert(self):
        self.driver.switch_to_alert()

    def back(self):
        self.driver.back()

    def forward(self):
        self.driver.forward()

    def iframe(self, id_index_locator):
        """常规切换 iframe，id_index_locator可以传入id，name ,locator"""
        try:
            if isinstance(id_index_locator, int):  # 如果传入的是数字，则以该数字为下标取值
                self.driver.switch_to.frame(id_index_locator)
            elif isinstance(id_index_locator, str):  # 如果传入的是字符串，则用iframe名字取值
                self.driver.switch_to.frame(id_index_locator)
            elif isinstance(id_index_locator, tuple):  # 如果是元祖，则根据传入的locator取值
                ele = self.find(id_index_locator)
                self.driver.switch_to.frame(ele)
        except:
            print("iframe切换异常")
            logger.info("iframe切换异常")
            self.screen()

    def alterMsg(self):
        """处理浏览器弹窗Alert窗口"""
        try:
            if EC.alert_is_present:
                Alert = self.driver.switch_to_alert()
                alertMsg = Alert.text
                logger.info('弹出的Alter信息：{}'.format(alertMsg))
                Alert.accept() #点击确定
            else:
                logger.info('没有弹出浏览器Alert窗口')
                alertMsg = '没有Alert窗口'
        except:
            logger.info("切换Alert异常")
            alertMsg = 'Alert窗口发生异常'
        return alertMsg


    def handle(self, value):
        """句柄切换，index、句柄名 """
        try:
            if isinstance(value, int):  # 切换到该下标对应的窗口
                handles = self.driver.window_handles
                self.driver.switch_to.window(handles[value])
            elif isinstance(value, str):  # 切换到该句柄名称对应的窗口
                self.driver.switch_to.window(value)
            else:
                print("传入的type参数 {value} 错误，仅可传int、str")
        except:
            print("根据 {value} 获取句柄失败")

    def quit_browse(self):
        self.driver.quit()
        time.sleep(1)
        sys.exit(1)

    def click_on_element(self, element):
        u"""鼠标悬停操作"""
        ActionChains(self.driver).click(element).perform()

    def find_element_click(self, locator):
        u"""鼠标悬停操作"""
        element = self.find(locator)
        ActionChains(self.driver).move_to_element(element).click().perform()

    def element_sendkey_click(self, locator,value):
        u"""鼠标悬停并用Enter代替点击"""
        element = self.find(locator)
        action = ActionChains(self.driver)
        action.send_keys_to_element(element,value)
        time.sleep(2)
        action.click(element)
        action.perform()

    def move_element_enter(self, locator):
        u"""鼠标悬停并键盘上输入Enter操作"""
        element = self.find(locator)
        ActionChains(self.driver).click(element).send_keys(Keys.ENTER).perform()
        time.sleep(1)

    def sendEnter(self):
        ActionChains(self.driver).send_keys(Keys.ENTER).perform()
        time.sleep(1)

    def get_text(self, locator):
        u"""获取文本内容"""
        return self.find_element(locator).text

    def get(self, locator, Type='text', name=''):
        ''' 根据传入的type判断获取指定的内容 （title、text、attribute）
        type==attribute: 获取元素属性 name:属性 className、name、text、value··· '''
        try:
            if Type == 'title':  # 获取当前网页 title
                return self.driver.title
            elif Type == 'text':  # 获取元素文本值（默认）
                return self.find(locator).text
            elif Type == 'attribute':  # 获取当前元素属性
                return self.find(locator).get_attribute(name)
            else:
                print("给的type参数 %s 错误，仅可用title、text、attribute",Type)
        except:
            print("获取 {Type} 值失败")
            return ''

    def get_attribute(self, locator, name):
        u"""获取属性"""
        return self.find(locator).get_attribute(name)

    def screen(self):
        u"""浏览器页面截图"""
        nowtime = time.strftime("%Y_%m_%d_%H_%M_%S")
        self.driver.get_screenshot_as_file(ReadConfig.get_image_path() + "%s.png" % nowtime)

    def get_screen_as_file(self, func):
        u"""异常自动截图"""
        def inner(*args, **kwargs):
            try:
                func(*args, **kwargs)
            except:
                self.screen()
                raise
        return inner

    """================处理Docx测试记录==============="""
    def screen_step(self,stepName):
        u"""浏览器页面截图"""
        nowtime = time.strftime("%Y%m%d%H%M%S")
        self.driver.get_screenshot_as_file(ReadConfig.get_image_path() + "%s.png" % nowtime)
        png = ReadConfig.get_image_path() + "%s.png" % nowtime
        self.doc.add_paragraph(stepName)
        self.doc.add_picture(png, width=Inches(6.0), height=Inches(3.0))

    def save_docreport(self,title):
        nowtime = time.strftime("%Y%m%d%H%M%S")
        title = title + '_' + nowtime
        self.doc.save(ReadConfig.get_reportDoc_path() + '%s.docx' %title)

    def add_dochead(self,title):
        return self.doc.add_heading(title,level=1)

    def screenshot_SaveAsDoc(self,stepName):
        self.screen_step(stepName)
        self.save_docreport(stepName)
        """================处理Docx测试记录==============="""

    def uploadFile(self,locator,filename):
        '''上传组件-上传完整文件
        :param locator: 上传按钮元素
        :param filename: 要上传的完整路径文件
        '''
        self.find_element_click(locator)
        time.sleep(2)
        args = r"E:\ProgramData\PycharmProjects\webtest\Common\upload.exe --chrome %s" % filename
        logger.info('要上传的文件路径:{}'.format(filename))
        print('执行上传的参数:{}'.format(args))
        os.system(args)
        time.sleep(3)


if __name__ == '__main__':
    print("用例开始执行时间：" + time.strftime("%Y%m%d%H%M%S"))
    print(ReadConfig.get_image_path())
    print(ReadConfig.get_reportDoc_path())










